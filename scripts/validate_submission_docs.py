"""Validate the generated submission documents without modifying them."""

from __future__ import annotations

import re
import zipfile
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
FORBIDDEN = (
    "\u56fd\u8d5b",
    "\u51b2\u523a",
    "\u7a0b\u5e8f\u5458\u4e3a\u672c",
    "AI\u5de5\u5177\u4e3a\u8f85\u52a9",
)


def main() -> None:
    output_dirs = [
        path
        for path in ROOT.iterdir()
        if path.is_dir() and list(path.glob("01_*.docx"))
    ]
    if len(output_dirs) != 1:
        raise RuntimeError(f"expected one submission-document directory: {output_dirs}")

    output_dir = output_dirs[0]
    for path in sorted(output_dir.iterdir()):
        if path.suffix.lower() == ".docx":
            validate_docx(path)
        elif path.suffix.lower() == ".txt":
            validate_script(path)


def validate_docx(path: Path) -> None:
    doc = Document(path)
    text_parts = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            text_parts.extend(cell.text for cell in row.cells)
    full_text = "\n".join(text_parts)
    forbidden_hits = [word for word in FORBIDDEN if word in full_text]

    with zipfile.ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    percent_tables = len(re.findall(r'w:type="pct"', document_xml))
    headings = sum(
        1
        for paragraph in doc.paragraphs
        if paragraph.style and paragraph.style.name.startswith("Heading")
    )
    list_paragraphs = sum(
        1 for paragraph in doc.paragraphs if paragraph._p.xpath("./w:pPr/w:numPr")
    )

    if forbidden_hits:
        raise AssertionError(f"{path.name}: forbidden terms {forbidden_hits}")
    if percent_tables:
        raise AssertionError(f"{path.name}: percent-width tables found")
    if not doc.tables or not headings:
        if not path.name.startswith("03_"):
            raise AssertionError(f"{path.name}: expected headings and tables")

    extra = ""
    if path.name.startswith("03_"):
        cjk_count = len(re.findall(r"[\u4e00-\u9fff]", full_text))
        timed_sections = len(
            re.findall(r"^\d{2}\s+\d:\d{2}—\d:\d{2}$", full_text, re.MULTILINE)
        )
        if timed_sections != 9:
            raise AssertionError(
                f"{path.name}: expected 9 timed sections, got {timed_sections}"
            )
        extra = (
            f", cjk={cjk_count}, timed_sections={timed_sections}, "
            f"estimated_minutes_at_240_cjk_per_min={cjk_count / 240:.2f}"
        )

    section = doc.sections[0]
    print(
        f"{path.name}: sections={len(doc.sections)}, "
        f"page_size={section.page_width}x{section.page_height}, "
        f"paragraphs={len(doc.paragraphs)}, headings={headings}, "
        f"tables={len(doc.tables)}, lists={list_paragraphs}, "
        f"percent_tables=0, forbidden_terms=0{extra}"
    )


def validate_script(path: Path) -> None:
    full_text = path.read_text(encoding="utf-8-sig")
    forbidden_hits = [word for word in FORBIDDEN if word in full_text]
    cjk_count = len(re.findall(r"[\u4e00-\u9fff]", full_text))
    nonspace_count = len(re.sub(r"\s+", "", full_text))
    timed_sections = len(
        re.findall(r"^【\d:\d{2}—\d:\d{2}】", full_text, re.MULTILINE)
    )

    if forbidden_hits:
        raise AssertionError(f"{path.name}: forbidden terms {forbidden_hits}")
    if timed_sections != 9:
        raise AssertionError(f"{path.name}: expected 9 timed sections, got {timed_sections}")

    print(
        f"{path.name}: cjk={cjk_count}, nonspace={nonspace_count}, "
        f"timed_sections={timed_sections}, "
        f"estimated_minutes_at_240_cjk_per_min={cjk_count / 240:.2f}, "
        "forbidden_terms=0"
    )


if __name__ == "__main__":
    main()

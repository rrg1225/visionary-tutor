from __future__ import annotations

import tempfile
import textwrap
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "提交材料"

FONT_CN = "Microsoft YaHei"
FONT_EN = "Aptos"
MONO = "Consolas"
INK = "20252B"
NAVY = "17324D"
BLUE = "2E74B5"
TEAL = "0F8F83"
MUTED = "5F6B76"
PALE_BLUE = "E8EEF5"
PALE_TEAL = "E8F5F3"
PALE_GRAY = "F2F4F7"
WHITE = "FFFFFF"
GREEN = "177245"
AMBER = "8A5A00"
RED = "9B1C1C"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def pil_color(value: str) -> str:
    return value if value.startswith("#") else f"#{value}"


def set_run(run, *, size: float | None = None, bold: bool | None = None,
            color: str | None = None, italic: bool | None = None,
            font: str = FONT_CN):
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), FONT_EN if font == FONT_CN else font)
    rpr.rFonts.set(qn("w:hAnsi"), FONT_EN if font == FONT_CN else font)
    rpr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = rgb(color)
    return run


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_borders(table, color="D7DEE5", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths: Sequence[int], indent=TABLE_INDENT_DXA):
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError(f"table widths must sum to {TABLE_WIDTH_DXA}: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, end))
    set_run(run, size=9, color=MUTED)


def add_numbering(doc: Document, ordered: bool) -> int:
    root = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in root.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in root.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    lvl.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "%1." if ordered else "•")
    lvl.append(text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    ppr.append(ind)
    lvl.append(ppr)
    abstract.append(lvl)
    # OOXML requires all abstract numbering definitions to precede concrete
    # <w:num> instances. Word may silently remap an out-of-order definition,
    # which can turn bullets into a continuing decimal list during PDF export.
    first_num = root.find(qn("w:num"))
    if first_num is None:
        root.append(abstract)
    else:
        root.insert(root.index(first_num), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    if ordered:
        # Word can continue numbering when two independent lists share the same
        # appearance. Force every logical ordered list to restart at one.
        level_override = OxmlElement("w:lvlOverride")
        level_override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), "1")
        level_override.append(start_override)
        num.append(level_override)
    root.append(num)
    return num_id


def configure_document(doc: Document, short_title: str):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = FONT_CN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = FONT_CN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = FONT_CN
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    caption.font.size = Pt(9)
    caption.font.color.rgb = rgb(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(header.add_run(short_title), size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(footer.add_run("VisionaryTutor  |  "), size=9, color=MUTED)
    add_page_number(footer)

    doc._bullet_num_id = add_numbering(doc, False)
    doc._decimal_num_id = add_numbering(doc, True)


def add_para(doc: Document, text: str = "", *, bold=False, color: str | None = None,
             size: float | None = None, align=None, before=0, after=6,
             italic=False, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.keep_with_next = keep
    if align is not None:
        p.alignment = align
    set_run(p.add_run(text), size=size, bold=bold, color=color, italic=italic)
    return p


def add_list(doc: Document, items: Iterable[str], ordered=False):
    num_id = add_numbering(doc, True) if ordered else doc._bullet_num_id
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.167
        ppr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num = OxmlElement("w:numId")
        num.set(qn("w:val"), str(num_id))
        num_pr.extend((ilvl, num))
        ppr.append(num_pr)
        set_run(p.add_run(item), size=10.6)


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]],
              widths: Sequence[int], *, header_fill=PALE_BLUE, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_header(table.rows[0])
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade_cell(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(str(value)), size=font_size, bold=True, color=NAVY)
    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            if len(str(value)) < 16 and index != len(row_values) - 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run(p.add_run(str(value)), size=font_size)
    set_table_geometry(table, widths)
    add_para(doc, "", after=2)
    return table


def add_callout(doc: Document, label: str, text: str, fill=PALE_TEAL, color=TEAL):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    set_table_borders(table, color="BFDAD6", size="5")
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(f"{label}  "), size=10.2, bold=True, color=color)
    set_run(p.add_run(text), size=10.2)
    add_para(doc, "", after=2)


def add_code_block(doc: Document, lines: Sequence[str]):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    set_table_borders(table, color="D5DAE0", size="4")
    cell = table.cell(0, 0)
    shade_cell(cell, "F6F7F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for index, line in enumerate(lines):
        if index:
            p.add_run().add_break()
        set_run(p.add_run(line), size=8.7, color="30363D", font=MONO)
    add_para(doc, "", after=2)


def add_cover(doc: Document, kicker: str, title: str, subtitle: str, metadata: Sequence[str]):
    add_para(doc, kicker, size=11, bold=True, color=TEAL,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=90, after=18)
    add_para(doc, title, size=28, bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_para(doc, subtitle, size=14, color=BLUE,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=44)
    for line in metadata:
        add_para(doc, line, size=10.5, color=MUTED,
                 align=WD_ALIGN_PARAGRAPH.CENTER, after=5)
    add_para(doc, "VisionaryTutor / 智眸学伴", size=12, bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=96, after=5)
    add_para(doc, "计算机视觉与深度学习个性化学习智能体", size=10, color=MUTED,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    doc.add_page_break()


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(text), size=9, color=MUTED)


def font_for(size: int, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def draw_centered(draw, box, text, font, fill, spacing=6):
    x1, y1, x2, y2 = box
    wrapped = textwrap.fill(text, width=max(6, int((x2 - x1) / (font.size * 1.05))))
    bbox = draw.multiline_textbbox((0, 0), wrapped, font=font, spacing=spacing, align="center")
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    draw.multiline_text(((x1 + x2 - width) / 2, (y1 + y2 - height) / 2), wrapped,
                        font=font, fill=pil_color(fill), spacing=spacing, align="center")


def arrow(draw, start, end, fill=BLUE, width=5):
    fill = pil_color(fill)
    draw.line((start, end), fill=fill, width=width)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        sign = 1 if ex > sx else -1
        points = [(ex, ey), (ex - sign * 18, ey - 10), (ex - sign * 18, ey + 10)]
    else:
        sign = 1 if ey > sy else -1
        points = [(ex, ey), (ex - 10, ey - sign * 18), (ex + 10, ey - sign * 18)]
    draw.polygon(points, fill=fill)


def make_architecture(path: Path):
    image = Image.new("RGB", (1600, 860), pil_color(WHITE))
    draw = ImageDraw.Draw(image)
    title = font_for(38, True)
    h = font_for(26, True)
    body = font_for(21)
    draw.text((70, 42), "VisionaryTutor 分层架构", font=title, fill=pil_color(NAVY))
    boxes = [
        (80, 145, 350, 690, PALE_TEAL, "体验层", "Vue 3 / Pinia\n流式对话\n画像、资源、路径\n实验与学习报告"),
        (390, 145, 660, 690, PALE_BLUE, "服务层", "Spring Boot 3\nJWT / JPA / Flyway\n会话、资源、推荐\n评估与审计 API"),
        (700, 145, 970, 690, "FFF5DF", "智能体层", "Planner / Supervisor\n8 类 Specialist\nCritic / Review\n共享黑板与交接"),
        (1010, 145, 1280, 690, "F3EAF8", "AI 与知识层", "LLM / RAG\nChroma + BM25\n引用与事实核验\n语音 / 视觉 / 视频接口"),
        (1320, 145, 1530, 690, PALE_GRAY, "数据层", "MySQL\nRedis\n课程语料\n对象存储\n审计日志"),
    ]
    for x1, y1, x2, y2, fill, label, content in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=24, fill=pil_color(fill), outline="#C9D3DD", width=3)
        draw_centered(draw, (x1 + 15, y1 + 30, x2 - 15, y1 + 100), label, h, NAVY)
        draw_centered(draw, (x1 + 18, y1 + 125, x2 - 18, y2 - 30), content, body, INK, 12)
    for left, right in zip(boxes, boxes[1:]):
        arrow(draw, (left[2] + 8, 420), (right[0] - 8, 420), TEAL, 5)
    draw.rounded_rectangle((190, 750, 1420, 820), radius=18, fill=pil_color(NAVY))
    draw_centered(draw, (210, 758, 1400, 812), "全链路：身份与权限隔离  ·  进度事件  ·  版本记录  ·  人工复核门禁", h, WHITE)
    image.save(path)


def make_agent_flow(path: Path):
    image = Image.new("RGB", (1600, 840), pil_color(WHITE))
    draw = ImageDraw.Draw(image)
    title = font_for(38, True)
    h = font_for(24, True)
    body = font_for(19)
    draw.text((70, 40), "多智能体协同与质量闭环", font=title, fill=pil_color(NAVY))
    top = [
        (70, 145, 300, 260, "学习画像 + 任务"),
        (360, 145, 590, 260, "Planner\n拆解与路由"),
        (650, 145, 930, 260, "Supervisor\n并行调度 / 共享黑板"),
        (1280, 145, 1530, 260, "发布资源\n路径与推荐"),
    ]
    for box in top:
        draw.rounded_rectangle(box[:4], radius=20, fill=pil_color(PALE_BLUE), outline=pil_color(BLUE), width=3)
        draw_centered(draw, box[:4], box[4], h, NAVY)
    arrow(draw, (300, 202), (360, 202), TEAL)
    arrow(draw, (590, 202), (650, 202), TEAL)
    specialists = ["讲义", "题库", "思维导图", "拓展阅读", "学习路径", "代码实操", "视频脚本", "交互可视化"]
    for i, label in enumerate(specialists):
        col = i % 4
        row = i // 4
        x1 = 160 + col * 300
        y1 = 355 + row * 150
        box = (x1, y1, x1 + 230, y1 + 90)
        draw.rounded_rectangle(box, radius=16, fill=pil_color(PALE_TEAL), outline=pil_color(TEAL), width=3)
        draw_centered(draw, box, label, h, NAVY)
        arrow(draw, (790, 260), (x1 + 115, y1), "7B8FA3", 3)
    critic = (1280, 380, 1530, 500)
    review = (1280, 575, 1530, 695)
    draw.rounded_rectangle(critic, radius=20, fill="#FFF5DF", outline=pil_color(AMBER), width=3)
    draw_centered(draw, critic, "Critic\n逐资源审查与返修", h, NAVY)
    draw.rounded_rectangle(review, radius=20, fill="#F3EAF8", outline="#7C4D99", width=3)
    draw_centered(draw, review, "Review\n最终质量门禁", h, NAVY)
    arrow(draw, (1200, 520), (1280, 440), AMBER)
    arrow(draw, (1405, 500), (1405, 575), "7C4D99")
    arrow(draw, (1405, 380), (1405, 270), TEAL)
    arrow(draw, (1405, 145), (1405, 110), TEAL)
    draw.text((1070, 760), "失败 → 有界返修 → 仍不合格则人工审核", font=body, fill=pil_color(RED))
    image.save(path)


def make_learning_loop(path: Path):
    image = Image.new("RGB", (1600, 560), pil_color(WHITE))
    draw = ImageDraw.Draw(image)
    title = font_for(38, True)
    h = font_for(22, True)
    body = font_for(18)
    draw.text((70, 38), "个性化学习闭环", font=title, fill=pil_color(NAVY))
    labels = [
        ("对话与测评", "抽取画像与知识状态"),
        ("诊断", "定位薄弱点和错误模式"),
        ("生成", "多智能体产出多模态资源"),
        ("学习", "阅读、练习、实验与答疑"),
        ("评估", "正确率、掌握度与行为证据"),
        ("调整", "更新画像、路径与推送策略"),
    ]
    xs = [70, 320, 570, 820, 1070, 1320]
    for i, ((label, detail), x) in enumerate(zip(labels, xs)):
        box = (x, 175, x + 205, 350)
        fill = PALE_TEAL if i % 2 == 0 else PALE_BLUE
        draw.rounded_rectangle(box, radius=22, fill=pil_color(fill), outline=pil_color(TEAL if i % 2 == 0 else BLUE), width=3)
        draw_centered(draw, (x + 10, 190, x + 195, 245), label, h, NAVY)
        draw_centered(draw, (x + 14, 255, x + 191, 335), detail, body, INK)
        if i < len(labels) - 1:
            arrow(draw, (x + 205, 262), (xs[i + 1] - 10, 262), TEAL, 4)
    draw.arc((160, 365, 1450, 520), start=0, end=180, fill=pil_color(TEAL), width=5)
    arrow(draw, (160, 442), (72, 350), TEAL, 5)
    draw.text((560, 430), "新的证据回流，学习方案随学随新", font=h, fill=pil_color(TEAL))
    image.save(path)


def add_picture(doc: Document, path: Path, caption: str, width=6.35):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def build_project_document(temp_dir: Path):
    architecture = temp_dir / "architecture.png"
    agent_flow = temp_dir / "agent_flow.png"
    learning_loop = temp_dir / "learning_loop.png"
    make_architecture(architecture)
    make_agent_flow(agent_flow)
    make_learning_loop(learning_loop)

    doc = Document()
    configure_document(doc, "项目开发与测试说明书")
    add_cover(
        doc,
        "项目技术文档",
        "项目开发与测试说明书",
        "基于大模型的个性化资源生成与学习多智能体系统",
        ("文档版本：V4.0", "编制日期：2026 年 7 月 18 日", "适用范围：作品审查、部署复现与演示说明"),
    )

    doc.add_heading("文档控制", level=1)
    add_table(doc, ("项目", "内容"), (
        ("系统名称", "VisionaryTutor / 智眸学伴"),
        ("示范课程", "计算机视觉与深度学习"),
        ("交付形态", "Web 应用，包含前端、后端、AI 引擎、课程知识库与测试包"),
        ("当前基线", "Git 提交 1e6046bd；Spring Boot 3.3.5；Java 17；Vue 3"),
        ("文档状态", "可提交版本"),
    ), (2100, 7260))
    add_callout(doc, "说明", "本文中的测试数字均来自 2026 年 7 月 18 日在当前代码基线上的实际执行结果；需要外部密钥、Docker 或真实云服务的测试边界单独列明。")

    doc.add_heading("摘要", level=1)
    add_para(doc, "VisionaryTutor 面向高校学生在资源筛选、学习节奏、知识短板识别和持续反馈方面的真实痛点，以计算机视觉与深度学习课程为示范场景，构建对话画像、多智能体资源生成、检索增强、动态学习路径、智能辅导和学习效果评估的一体化学习闭环。系统不是单次问答页面，而是把学生画像、知识证据、生成资源、学习行为和评估结果纳入同一会话与审计链路。")
    add_para(doc, "系统可生成讲义、练习题、思维导图、学习路径、代码实操、拓展阅读、视频脚本和交互可视化八类资源。Planner、Supervisor、专用生成智能体、Critic 与 Review 共同完成任务拆解、并行生成、逐项复核、有界返修和发布门禁；RAG 引用验证、检索内容隔离、事实检查与人工复核共同降低幻觉风险。")

    doc.add_heading("目录", level=1)
    add_table(doc, ("章节", "内容"), (
        ("1", "项目背景、用户痛点与建设目标"),
        ("2", "需求分析与功能覆盖"),
        ("3", "总体架构与技术选型"),
        ("4", "对话式动态学习画像"),
        ("5", "多智能体协同与资源生成"),
        ("6", "知识库、RAG 与防幻觉机制"),
        ("7", "学习路径、资源推送、辅导与评估"),
        ("8", "交互设计与多模态体验"),
        ("9", "数据、安全与可审计设计"),
        ("10", "部署、配置与运行维护"),
        ("11", "测试策略、结果与边界"),
        ("12", "性能、可靠性与降级"),
        ("13", "创新价值与应用前景"),
        ("14", "开源组件、AI 工具与交付清单"),
        ("附录", "复现命令与验收检查表"),
    ), (1200, 8160))

    doc.add_heading("1 项目背景、用户痛点与建设目标", level=1)
    doc.add_heading("1.1 用户与场景", level=2)
    add_para(doc, "目标用户为正在学习人工智能、计算机视觉、深度学习及相关课程的高校学生。典型场景包括课前补基础、课中理解抽象概念、课后针对薄弱点练习、项目式实操、阶段复盘和考试准备。教师与课程建设者可通过管理侧能力查看知识库状态、内容审核与生成审计。")
    doc.add_heading("1.2 核心痛点", level=2)
    add_list(doc, (
        "学习资源数量多、来源分散，学生难以判断内容是否适合当前基础与目标。",
        "统一讲授无法持续适配不同学生的认知风格、节奏、易错点和掌握程度。",
        "普通问答工具缺少课程证据、资源结构和学习路径，回答难以沉淀为可执行计划。",
        "多模态内容生成等待时间长，缺少进度提示、失败解释和可恢复机制。",
        "生成内容可能存在事实、引用、安全或授权风险，需要系统化质量门禁。",
    ))
    doc.add_heading("1.3 建设目标", level=2)
    add_list(doc, (
        "以自然对话构建不少于六个维度的动态画像，并使用测验与行为证据持续校准。",
        "通过多智能体协作生成至少五类、实际支持八类个性化学习资源。",
        "将画像、知识状态、学习进度和资源使用反馈转化为动态路径与精准推送。",
        "提供流式辅导、交互实验、代码沙箱、学习报告与可追溯质量审查。",
        "形成可部署、可测试、可降级、可审计且尊重语料授权的工程化系统。",
    ))

    doc.add_heading("2 需求分析与功能覆盖", level=1)
    add_table(doc, ("需求域", "系统实现", "验收观察点"), (
        ("对话画像", "七维画像、证据与置信度、知识状态、版本更新", "完成对话后可查看画像；练习结果会校准薄弱点"),
        ("资源生成", "八种 ArtifactType；多智能体并行生成与质量复核", "同一主题生成多类型资源，界面显示角色与进度"),
        ("学习路径", "DAG 路径节点、依赖关系、预计时长、失败回退", "路径按掌握度调整，未掌握节点触发补强"),
        ("资源推送", "画像、查询、学习状态与历史行为的混合推荐", "推荐卡片包含来源、原因与资源类型"),
        ("智能辅导", "流式 Markdown、上下文会话、引用、TTS 与图解资源", "回答可连续显示并保留上下文"),
        ("效果评估", "固定测验、练习记录、知识追踪、前后测与学习报告", "报告展示掌握变化与下一步建议"),
        ("内容安全", "检索隔离、引用校验、事实检查、Critic、人工审核", "低质量内容不会直接发布"),
        ("响应体验", "SSE/进度事件、任务状态、超时、重试和降级提示", "长任务无白屏，失败原因可见"),
    ), (1500, 4400, 3460), font_size=8.6)

    doc.add_heading("3 总体架构与技术选型", level=1)
    add_picture(doc, architecture, "图 1  系统总体分层架构")
    doc.add_heading("3.1 分层职责", level=2)
    add_para(doc, "前端负责学习工作台、资源卡片、进度呈现和交互实验；Spring Boot 后端承载身份、会话、画像、编排、资源、路径、推荐、评估和审计；Python AI 引擎负责课程文档处理、RAG 评估与演示文档能力；MySQL 保存长期业务状态，Redis 支持缓存和可选消息总线，Chroma 提供向量检索，BM25 作为高可用降级路径。")
    add_table(doc, ("层级", "主要技术", "选择理由"), (
        ("前端", "Vue 3、Vite、Pinia、Element Plus", "组件化、响应式状态与快速构建，适合学习工作台"),
        ("可视化", "Mermaid、KaTeX、MediaPipe、Pyodide", "支持图解、公式、本地视觉与浏览器内代码实验"),
        ("后端", "Spring Boot 3、Java 17、Security、JPA、Flyway", "成熟的分层服务、鉴权、事务和数据库演进能力"),
        ("智能体", "自研 Supervisor + Specialist + Critic 协议", "角色边界清晰，可追踪、可返修、可人工接管"),
        ("RAG", "LangChain4j、Chroma、BM25、RRF", "兼顾语义召回、词项精确度和故障降级"),
        ("测试", "JUnit、Vitest、Playwright、Pytest", "覆盖单元、组件、端到端与 AI 数据处理"),
    ), (1450, 3350, 4560), font_size=8.8)

    doc.add_heading("4 对话式动态学习画像", level=1)
    doc.add_heading("4.1 七个核心维度", level=2)
    add_table(doc, ("维度", "含义", "主要证据"), (
        ("knowledgeBase", "已有知识基础与前置概念掌握情况", "自述、测验、知识追踪"),
        ("goal", "学习目标、用途和阶段期望", "对话与任务选择"),
        ("cognitiveStyle", "偏好的解释、图示、示例与推导方式", "对话、资源反馈"),
        ("weakPoints", "待加强知识点", "错题、掌握度、行为指标"),
        ("errorPatterns", "重复出现的错误类型和推理偏差", "答题与代码运行记录"),
        ("learningPace", "适合的内容粒度、节奏与预计时长", "学习时长、完成情况"),
        ("emotionAttention", "学习状态与注意力相关辅助信号", "用户授权后的状态输入"),
    ), (1900, 3600, 3860), font_size=8.8)
    doc.add_heading("4.2 更新与冲突处理", level=2)
    add_para(doc, "每个画像维度保存 value、evidence 和 confidence。新对话先与既有快照合并，再由测验、资源使用和知识追踪指标校准；变化快的知识状态采用更快衰减，认知风格与学习节奏采用相对稳定的更新策略。画像更新形成版本和事件记录，避免覆盖后无法解释。")
    add_callout(doc, "画像原则", "画像是带证据、可更新的学习假设，不是对学生的永久标签。系统保留“待观察”状态，并允许后续证据修正。")

    doc.add_heading("5 多智能体协同与资源生成", level=1)
    add_picture(doc, agent_flow, "图 2  多智能体协同、返修与发布门禁")
    doc.add_heading("5.1 角色分工", level=2)
    add_table(doc, ("角色", "责任", "主要输出"), (
        ("Planner", "解析任务、画像与短板，制定执行计划", "目标、资源类型、约束和顺序"),
        ("Supervisor", "路由、并行调度、共享上下文、超时与状态控制", "可审计的 run 与步骤事件"),
        ("Doc / Quiz / MindMap", "分别生成讲义、题库与知识结构", "结构化资源内容"),
        ("Reading / Coding / Visualization", "扩展阅读、代码实践与交互可视化", "多模态学习材料"),
        ("Path", "综合画像与其他资源摘要生成学习路径", "DAG 节点、依赖和时长"),
        ("Critic", "逐资源检查事实性、结构、适配度与安全", "通过或返修意见"),
        ("Review", "汇总最终报告并执行发布前门禁", "质量结论与审计信息"),
    ), (1700, 3900, 3760), font_size=8.6)
    doc.add_heading("5.2 协作协议", level=2)
    add_list(doc, (
        "Planner 创建计划，Supervisor 为同一任务分配 runId 并建立共享黑板。",
        "专用智能体可先交换大纲摘要，再并行生成最终内容；Path 在需要时读取其他资源摘要。",
        "Critic 对每项产物单独评分并给出返修指令，返修次数由配置限制。",
        "达到返修上限仍不合格时标记 MANUAL_REVIEW_REQUIRED，不将风险内容伪装为成功。",
        "Review 汇总资源、引用、Critic 报告和运行元数据，满足门禁后发布。",
    ), ordered=True)
    doc.add_heading("5.3 支持的八类资源", level=2)
    add_table(doc, ("类型", "ArtifactType", "个性化要素"), (
        ("课程讲义", "HANDOUT", "基础层级、推导深度、示例密度"),
        ("练习题", "QUIZ", "难度、题型、薄弱点与解析方式"),
        ("思维导图", "MINDMAP", "知识依赖和重点分支"),
        ("学习路径", "LEARNING_PATH", "先后关系、时长、掌握门槛"),
        ("代码实操", "CODE_PRACTICE", "语言、脚手架、测试与运行提示"),
        ("拓展阅读", "EXTENDED_READING", "先修知识、阅读目标与延伸方向"),
        ("视频脚本", "VIDEO_SCRIPT", "讲解节奏、分段和画面要点"),
        ("交互可视化", "VISUALIZATION", "参数、动态过程与结果解释"),
    ), (1750, 2200, 5410), font_size=8.8)

    doc.add_heading("6 知识库、RAG 与防幻觉机制", level=1)
    doc.add_heading("6.1 知识数据生命周期", level=2)
    add_para(doc, "项目保留 493 份清洗后的课程与技术资料，以及来源、许可、哈希和提交资格元数据。大规模分块、向量索引和运行缓存不随源码提交，由 document_processor.py 从 cleaned 语料重建。数据集清单采用默认拒绝分发策略：授权未核清或具有商业版权风险的文档不会进入提交包。")
    doc.add_heading("6.2 检索与引用", level=2)
    add_list(doc, (
        "向量检索负责语义相关性，BM25 负责关键词精确匹配，RRF 用于融合候选结果。",
        "Chroma 不可用或超时时自动切换 BM25 高可用模式，并在响应元数据中记录降级。",
        "资源保存 citationsJson、validationStatus 与 verificationAuditJson，支持逐项查看依据。",
        "CitationValidator、GroundingAuditService 与 StrictFactCheckService 对引用与事实一致性进行检查。",
    ))
    doc.add_heading("6.3 内容安全与发布门禁", level=2)
    add_table(doc, ("风险", "控制措施", "失败处理"), (
        ("检索内容提示注入", "RetrievedContentGuard 隔离指令式片段", "隔离片段并记录"),
        ("事实无依据", "RAG grounding、引用验证、严格事实检查", "阻止发布或转人工"),
        ("结构不合规", "JSON Schema 与资源契约校验", "自动修复或重新生成"),
        ("不当内容", "模型安全策略、关键词规则、Critic 审查", "拒绝、降级或人工审核"),
        ("授权不清", "manifest 提交资格与哈希清单", "默认不分发"),
    ), (1800, 4450, 3110), font_size=8.7)

    doc.add_heading("7 学习路径、资源推送、辅导与评估", level=1)
    add_picture(doc, learning_loop, "图 3  学习证据驱动的动态闭环")
    doc.add_heading("7.1 动态路径", level=2)
    add_para(doc, "学习路径以 DAG 表达前置关系、补强关系与顺序，节点记录资源类型、预计学习时间和掌握度。某节点失败时，服务会回溯相关前置节点、安排强化资源并暂时阻塞后续依赖，避免学生在基础未掌握时被机械推向更难内容。")
    doc.add_heading("7.2 精准推荐", level=2)
    add_para(doc, "推荐服务结合用户画像、当前会话主题、薄弱点、资源类型偏好、历史使用记录和可用资源计算候选项，并记录推荐原因、来源和是否使用降级策略。资源卡片明确标识系统资源、个人生成资源或示例内容，减少来源混淆。")
    doc.add_heading("7.3 即时辅导", level=2)
    add_para(doc, "辅导面板支持流式 Markdown、公式、引用与上下文会话；可将讲解转换为语音，并联动思维导图、可视化或代码练习。对需要精确课程依据的问题优先使用 RAG；证据不足时给出边界提示而不是编造结论。")
    doc.add_heading("7.4 学习效果评估", level=2)
    add_para(doc, "系统记录固定测验、生成练习、错题复习、学习时长、资源反馈和知识点正确率，生成学习状态报告与前后测对比。评估结果回流画像、路径与推荐，并保留证据 ID，便于解释“为什么发生这次调整”。")

    doc.add_heading("8 交互设计与多模态体验", level=1)
    add_table(doc, ("界面能力", "实现要点", "用户价值"), (
        ("流式工作台", "逐段显示回答、状态与引用", "缩短等待感并保持上下文"),
        ("资源卡片", "类型化渲染、来源标识、状态与操作集中", "快速识别资源用途和可信状态"),
        ("生成进度", "规划、检索、生成、审查、持久化事件", "长任务始终有反馈"),
        ("CNN 互动实验", "可调 padding/stride、滑窗步骤和输出矩阵", "把抽象卷积过程变成可操作验证"),
        ("代码沙箱", "浏览器端 Pyodide 与后端隔离执行策略", "修改代码、运行并编写测试"),
        ("响应式设计", "桌面与移动端导航、卡片和阅读节奏", "适应多设备学习"),
        ("Markdown 安全", "markdown-it + DOMPurify + 受控组件", "兼顾表现力与 XSS 防护"),
    ), (1650, 4300, 3410), font_size=8.7)

    doc.add_heading("9 数据、安全与可审计设计", level=1)
    doc.add_heading("9.1 数据模型", level=2)
    add_para(doc, "当前后端包含 30 个 Flyway 数据库迁移。核心实体覆盖用户、学习会话、画像快照、学习路径节点与边、生成资源、生成事件、智能体执行日志、Critic 返修记录、测验与错题、知识追踪、资源使用、推荐日志、聊天消息、用户记忆、共享教材与内容审核。")
    doc.add_heading("9.2 身份与数据隔离", level=2)
    add_list(doc, (
        "Spring Security 与 JWT 负责认证，控制器和服务层同时校验资源所有权。",
        "游客会话与注册用户数据分离，转正时执行显式迁移，避免跨用户串读。",
        "真实密钥只从环境变量或本地配置注入；示例文件不包含有效凭据。",
        "密码重置对未知邮箱返回通用结果，避免账户枚举；输入使用 Bean Validation 校验。",
        "代码执行采用白名单、超时、资源限制与可选 Docker 隔离，不直接信任生成代码。",
    ))
    doc.add_heading("9.3 审计", level=2)
    add_para(doc, "每次资源生成记录 traceId、runId、模型、Prompt 版本、Agent、状态迁移、耗时、降级原因、引用审计和返修信息。审计页面可按运行追踪各智能体步骤，为问题定位、质量复盘和人工审核提供依据。")

    doc.add_heading("10 部署、配置与运行维护", level=1)
    doc.add_heading("10.1 环境要求", level=2)
    add_table(doc, ("组件", "建议版本", "用途"), (
        ("JDK", "17", "Spring Boot 后端"),
        ("Node.js", "20 及以上", "前端构建与开发"),
        ("Python", "3.11", "知识处理与评估"),
        ("MySQL", "8.x", "业务数据"),
        ("Redis", "7.x", "缓存、会话与可选消息总线"),
        ("Chroma", "0.5.x 接口兼容", "向量知识库"),
        ("Docker", "可选", "依赖编排、集成测试与代码隔离"),
    ), (1800, 2600, 4960), font_size=8.8)
    doc.add_heading("10.2 快速启动", level=2)
    add_code_block(doc, (
        "Copy-Item .env.example backend\\.env.properties",
        ".\\mvnw.cmd -f backend\\pom.xml -DskipTests package",
        "java -jar backend\\target\\visionary-tutor-backend-0.0.1-SNAPSHOT.jar",
        "cd frontend",
        "npm ci",
        "npm run dev",
    ))
    add_para(doc, "也可在项目根目录运行 scripts/start-all.ps1。首次启用向量知识库时，按 ai_engine/requirements.txt 安装依赖并运行 document_processor.py 重建分块与索引。生产配置应使用高强度 JWT 密钥、独立数据库账号、HTTPS、受限跨域来源和最小权限云服务凭据。")
    doc.add_heading("10.3 演示模式与降级", level=2)
    add_para(doc, "demo Profile 可初始化 CNN 示例场景。普通生产配置不默认创建演示账号，也不开放远程种子接口。模型密钥缺失时系统使用可解释的本地模板；Chroma 故障时切换 BM25；媒体生成超时保留文本脚本与已完成资源，并在界面显示状态。")

    doc.add_heading("11 测试策略、结果与边界", level=1)
    doc.add_heading("11.1 测试分层", level=2)
    add_para(doc, "测试覆盖领域逻辑、服务、控制器、安全、数据库仓储、智能体协议、RAG、资源契约、前端领域函数、组件渲染、响应式界面和浏览器业务流程。外部模型与真实基础设施通过 Mock、H2、可选 Testcontainers 和显式集成 Profile 分层处理。")
    doc.add_heading("11.2 本次实测结果", level=2)
    add_table(doc, ("测试项", "命令/范围", "结果"), (
        ("后端测试", "mvnw -f backend/pom.xml test", "208 项；0 失败；0 错误；1 项外部 MCP 测试跳过"),
        ("前端单元/组件", "npm run test:unit", "14 个文件、50 项全部通过"),
        ("前端覆盖率", "Vitest + V8", "语句 96.15%；分支 82.14%；函数 81.33%；行 96.15%"),
        ("Python 测试", "python -m pytest tests -q", "29 项全部通过"),
        ("浏览器 E2E", "npm run test:e2e", "19 项通过；1 项真实后端场景按配置跳过"),
        ("生产构建", "npm run build", "3315 模块构建成功；包体预算通过"),
        ("首屏 JS 预算", "check-bundle-budget.mjs", "612.1 KiB；Mermaid 保持懒加载"),
    ), (1700, 3300, 4360), font_size=8.7)
    add_callout(doc, "测试边界", "本次默认测试未声明真实云模型、真实 MySQL/Redis/Chroma 集成和视频生成服务全部可用。真实后端 E2E、Testcontainers 集成与云服务验收需要按部署环境配置密钥和 Docker 后单独执行。", fill="FFF5DF", color=AMBER)
    doc.add_heading("11.3 缺陷闭环", level=2)
    add_para(doc, "首次端到端执行发现一处测试仍依赖旧版空状态文案。经核对，页面功能正常而断言已过期；测试改为验证历史搜索控件在折叠状态下不存在，随后完整 E2E 复跑通过。该调整降低了对易变文案的耦合，并保留对真实交互状态的验证。")

    doc.add_heading("12 性能、可靠性与降级", level=1)
    add_table(doc, ("机制", "实现", "效果"), (
        ("并行生成", "专用智能体并行调度，路径智能体按依赖汇总", "减少多资源串行等待"),
        ("流式输出", "SSE/进度事件分阶段呈现", "避免长时间白屏"),
        ("有界超时", "模型、工具、智能体和总任务分别配置超时", "防止无限占用"),
        ("幂等与状态机", "runId、生成状态和持久化约束", "刷新或重试不重复生成"),
        ("检索降级", "Chroma → BM25", "向量服务异常时仍可提供课程检索"),
        ("模型降级", "云模型 → 本地模板/已有资源", "缺少密钥或超时仍可解释"),
        ("质量熔断", "Critic 返修上限与人工审核", "避免低质量无限重试或直接发布"),
        ("可观测性", "健康检查、Prometheus、追踪与审计事件", "快速定位性能和质量问题"),
    ), (1750, 4400, 3210), font_size=8.7)

    doc.add_heading("13 创新价值与应用前景", level=1)
    doc.add_heading("13.1 技术创新", level=2)
    add_list(doc, (
        "把对话画像、知识追踪和行为指标融合为带置信度、可衰减、可校准的动态画像。",
        "以可审计的 Supervisor 协议组织多角色生成，并加入逐资源 Critic、有界返修和人工审核门禁。",
        "将多模态资源放入同一学习路径，而非作为彼此孤立的生成结果。",
        "检索、引用、授权、内容安全和发布状态贯穿数据生命周期，形成工程化防幻觉体系。",
        "CNN 可视化与代码验证共享学习主题，形成“观察—解释—实现—验证”的实践闭环。",
    ))
    doc.add_heading("13.2 实用价值", level=2)
    add_para(doc, "对学生，系统降低资源筛选成本并提供连续的个性化反馈；对教师，系统可承担资源初稿、分层练习和学习证据汇总，教师保留内容审核与教学决策；对学校，系统可扩展到更多课程，沉淀可追踪的课程知识库、资源资产和学习数据治理能力。")
    doc.add_heading("13.3 扩展路线", level=2)
    add_para(doc, "系统的 ArtifactType、智能体注册、知识层、课程数据与前端资源卡片均采用可扩展设计，可进一步支持电子信息、数据结构、信号处理等课程。后续可在获得授权与伦理审查后开展真实用户研究、A/B 对照和长期学习效果验证。")

    doc.add_heading("14 开源组件、AI 工具与交付清单", level=1)
    doc.add_heading("14.1 开源与云服务", level=2)
    add_para(doc, "第三方组件名称、来源和协议集中记录在 THIRD_PARTY.md；逐文档语料授权、SHA-256 和提交资格记录在 datasets/manifest.json。云模型与语音、视觉、视频服务按各自服务协议配置，API Key 不进入仓库。")
    doc.add_heading("14.2 AI Coding 工具", level=2)
    add_para(doc, "开发过程中使用 Claude Code、OpenAI Codex 等编码辅助工具处理代码检索、重复性脚手架、测试用例草稿、构建问题定位和文档整理。需求定义、架构边界、业务规则、最终实现选择、代码审核、测试验收和提交责任由团队成员承担。详细边界见《AI Coding 工具使用说明》。")
    doc.add_heading("14.3 交付清单", level=2)
    add_table(doc, ("目录/文件", "内容", "是否必要"), (
        ("backend/", "Spring Boot 源码、配置、迁移与测试", "是"),
        ("frontend/", "Vue 源码、锁文件、单元测试与配置", "是"),
        ("ai_engine/", "知识处理、评估、核心语料与测试", "是"),
        ("datasets/", "数据来源、授权与哈希清单", "是"),
        ("tests/、loadtest/", "E2E 与 JMeter 测试", "是"),
        ("scripts/", "启动、校验、打包与文档生成脚本", "是"),
        ("提交材料/", "本说明书、AI 使用说明、演示讲稿", "是"),
        ("THIRD_PARTY.md", "第三方组件与协议", "是"),
        ("node_modules、target、dist、output", "可再生依赖、构建与运行产物", "不提交"),
    ), (2200, 5200, 1960), font_size=8.8)

    doc.add_heading("附录 A 复现与质量检查命令", level=1)
    add_code_block(doc, (
        ".\\mvnw.cmd -f backend\\pom.xml test",
        "cd frontend",
        "npm ci",
        "npm run test:unit",
        "npm run build",
        "npm run test:e2e",
        "cd ..\\ai_engine",
        "python -m pip install -r requirements.txt",
        "python -m pytest tests -q",
        "cd ..",
        ".\\scripts\\build-submission.ps1",
    ))
    doc.add_heading("附录 B 提交前验收检查表", level=1)
    add_list(doc, (
        "确认 .env.example 仅为示例，提交包中不存在真实 API Key、数据库密码和私有地址。",
        "确认 backend、frontend、ai_engine、tests、datasets、scripts 和依赖锁文件完整。",
        "确认 node_modules、target、dist、coverage、test-results、运行日志和向量缓存未进入提交包。",
        "确认 THIRD_PARTY.md 与 datasets/manifest.json 随包提供，授权不明语料未分发。",
        "按本说明书复跑后端、前端、Python 和 E2E 测试，记录环境差异与跳过原因。",
        "演示前检查 demo Profile、模型密钥、Chroma 状态和降级资源；不在现场暴露密钥。",
        "核对演示视频时长、字幕清晰度、生成进度、引用、审计和学习闭环均可见。",
    ))

    doc.core_properties.title = "VisionaryTutor 项目开发与测试说明书"
    doc.core_properties.subject = "个性化学习多智能体系统"
    doc.core_properties.author = "VisionaryTutor 项目团队"
    doc.core_properties.keywords = "多智能体, 个性化学习, RAG, 资源生成, 学习评估"
    doc.save(OUT / "01_项目开发与测试说明书.docx")


def build_ai_coding_document():
    doc = Document()
    configure_document(doc, "AI Coding 工具使用说明")
    add_cover(
        doc,
        "开发过程说明",
        "AI Coding 工具使用说明",
        "工具用途、工作边界、人工审核与责任声明",
        ("文档版本：V1.0", "编制日期：2026 年 7 月 18 日", "适用项目：VisionaryTutor / 智眸学伴"),
    )

    doc.add_heading("1 说明目的", level=1)
    add_para(doc, "本说明用于如实披露项目开发过程中 AI Coding 工具的使用方式，明确工具参与的任务、未被委托的关键决策、人工审核流程及安全边界。AI 生成或建议的内容均作为候选草稿进入工程流程，只有在团队成员理解、修改、验证并确认后才可保留。")
    add_callout(doc, "责任边界", "工具不拥有需求解释权、架构决策权、代码合并权或发布权。团队成员对最终提交的源码、测试、配置、数据授权和文档内容负责。")

    doc.add_heading("2 使用的工具与主要用途", level=1)
    add_table(doc, ("工具", "使用场景", "典型产出", "处理方式"), (
        ("Claude Code", "仓库检索、局部实现建议、重复代码与测试草稿", "候选补丁、测试结构、问题定位线索", "逐行审核后修改或弃用"),
        ("OpenAI Codex", "代码审阅、测试补充、构建排查、文档整理", "审阅意见、候选测试、命令与文档初稿", "结合源码和实际执行结果复核"),
        ("运行时大模型服务", "系统画像、资源生成、辅导与质量审查", "面向学习者的内容", "经 RAG、规则、Critic 与人工门禁"),
    ), (1600, 3000, 2800, 1960), font_size=8.6)
    add_para(doc, "编码辅助工具主要用于减少机械性工作，包括查找调用关系、生成 DTO/测试脚手架、补齐边界用例、整理重复配置、解释构建错误和统一文档格式。对于需要上下文判断的业务逻辑，工具建议只作为讨论材料。")

    doc.add_heading("3 工具参与与人工负责事项", level=1)
    add_table(doc, ("工作内容", "工具可参与", "团队成员必须完成"), (
        ("需求分析", "归纳需求、列出验收点", "确认用户问题、优先级和真实边界"),
        ("架构设计", "比较方案、提示常见风险", "确定模块边界、数据流、故障策略和技术债"),
        ("编码", "生成局部草稿、重复样板和重构建议", "理解代码、调整实现、处理并发/事务/权限"),
        ("测试", "提出用例、生成 Mock 和断言草稿", "检查断言有效性、复现缺陷、执行完整测试"),
        ("安全", "提示常见漏洞与校验项", "审查鉴权、数据隔离、密钥、输入和执行沙箱"),
        ("文档", "整理结构、润色和格式化", "核对每项功能、数字、协议和测试边界"),
        ("提交", "提供清单和命令建议", "决定最终内容、完成验收并承担责任"),
    ), (1850, 3350, 4160), font_size=8.7)

    doc.add_heading("4 标准使用流程", level=1)
    add_list(doc, (
        "团队成员先给出明确任务、输入范围、验收条件和禁止事项。",
        "工具读取限定范围内的代码或错误信息，返回解释、候选补丁或测试建议。",
        "团队成员逐行检查逻辑，重点核对权限、事务、异常、并发、资源释放和数据边界。",
        "候选代码经过格式检查、静态分析、单元测试、组件测试或端到端测试；测试失败必须定位原因。",
        "涉及模型输出、语料授权、用户数据或安全策略的变更需要额外人工复核。",
        "只有通过审核与验证的修改才进入项目基线；无法解释或无法验证的建议不采用。",
    ), ordered=True)

    doc.add_heading("5 测试包中的 AI 辅助范围", level=1)
    add_para(doc, "AI Coding 工具参与了部分测试用例草稿、Mock 数据结构、边界条件枚举和重复断言整理。例如，针对权限所有权、输入校验、生成状态机、RAG 降级、资源契约、前端卡片渲染和响应式页面，工具可先生成候选测试结构。团队成员随后根据真实接口和实现修改用例，并实际执行测试。")
    add_table(doc, ("质量要求", "人工审核要点"), (
        ("断言有效", "测试必须验证业务结果或状态，不只验证函数被调用"),
        ("与实现解耦", "优先验证外部行为，避免绑定易变文案和内部细节"),
        ("覆盖异常", "包含超时、无权限、空数据、重复请求、外部服务失败与降级"),
        ("数据隔离", "不同用户、游客与注册用户场景不得互相读取资源"),
        ("可重复", "不依赖随机外部响应；真实服务测试由独立 Profile 控制"),
        ("失败可解释", "跳过项和外部依赖边界必须写明，不能把未执行写成通过"),
    ), (2300, 7060), font_size=9.0)
    add_para(doc, "当前材料引用的测试数字来自实际命令输出：后端 208 项测试通过且 1 项外部 MCP 场景跳过；前端 50 项单元/组件测试通过；Python 29 项测试通过；浏览器 E2E 19 项通过、1 项真实后端场景按配置跳过。")

    doc.add_heading("6 代码审核重点", level=1)
    add_list(doc, (
        "业务正确性：输入、状态转换、幂等、事务边界和失败恢复是否符合需求。",
        "权限安全：JWT、资源所有权、管理员权限、游客迁移和跨用户数据隔离。",
        "隐私与密钥：日志不输出敏感值，真实凭据不写入代码、提示词或提交文件。",
        "模型边界：Prompt 注入、无依据回答、结构化输出校验、返修上限和人工审核。",
        "工程质量：命名、可读性、异常处理、超时、资源释放、兼容性和可测试性。",
        "依赖与许可：新增组件和语料必须确认来源、版本、协议及再分发条件。",
    ))

    doc.add_heading("7 数据、安全与保密约束", level=1)
    add_para(doc, "向编码辅助工具提供上下文时，应遵循最小必要原则，不粘贴真实 API Key、密码、个人身份信息、未公开商业数据或授权不明的完整资料。项目密钥仅存在于本地环境变量或不提交的配置文件。若工具输出包含疑似密钥、真实个人信息或不明来源代码，应立即停止使用并进行清理与溯源。")
    add_para(doc, "AI 建议的第三方代码不能直接复制进入项目。团队成员需要核查许可证、原始来源和兼容性；无法确认来源时，应自行实现或选择许可清晰的依赖。")

    doc.add_heading("8 局限与处理原则", level=1)
    add_list(doc, (
        "工具可能误读跨模块上下文，因此所有建议必须回到源码、接口和测试验证。",
        "工具可能生成不存在的 API、配置或测试结果；文档只采用实际代码和实际执行证据。",
        "工具可能倾向一次性大改；项目采用小范围变更、明确回滚点和分层测试。",
        "模型输出不稳定；关键规则使用确定性代码、契约校验和人工审批固定下来。",
    ))

    doc.add_heading("9 最终声明", level=1)
    add_para(doc, "本项目使用 AI Coding 工具提高信息检索、重复性编码、测试设计和文档整理效率。工具没有替代团队成员对需求、架构、业务规则、代码质量、安全、授权和测试结论的判断。所有保留的代码与材料均应达到可解释、可审核、可复现的要求；团队成员对最终交付承担完整责任。", bold=True, color=NAVY, size=11.2)
    add_para(doc, "团队签字：____________________    日期：____________________", before=28, after=8)

    doc.core_properties.title = "VisionaryTutor AI Coding 工具使用说明"
    doc.core_properties.subject = "AI 编码辅助使用边界与人工审核"
    doc.core_properties.author = "VisionaryTutor 项目团队"
    doc.save(OUT / "02_AI_Coding工具使用说明.docx")


def build_video_script_document():
    script = """【0:00—0:40】
各位评审老师好，我们的项目是 VisionaryTutor，中文名“智眸学伴”。它面向高校计算机视觉与深度学习课程，解决三个常见问题：学习资源很多，却很难找到适合自己的内容；统一课程节奏难以照顾每个人的基础和薄弱点；普通问答给出答案后，往往没有后续练习、学习路径和效果反馈。我们的目标，是让系统真正理解学生当前处于什么水平、下一步应该学什么，并持续生成能够直接学习和实践的个性化资源。整套体验围绕一个原则展开：每一项推荐和调整，都应有学习证据，也应让学生看得懂原因。

【0:40—1:30】
学生第一次进入系统时，不需要填写冗长表单，而是通过自然语言说明自己的专业、目标、学习经历和困惑。系统会从对话中形成七个维度的学习画像，包括知识基础、学习目标、认知风格、薄弱点、错误模式、学习节奏，以及学习状态辅助信号。每个维度都带有证据和置信度，不确定的信息会保留为“待观察”，不会直接给学生贴标签。学生可以查看画像依据，也可以纠正系统理解有偏差的部分。随着学生完成练习、阅读资源和参加测验，系统会把新的正确率、掌握度和行为证据写回画像，实现随学随新，而不是让首次对话的结论长期固定。

【1:30—2:35】
接下来展示多智能体资源生成。学生输入“我总是弄不清 padding、stride 和卷积输出尺寸的关系，希望通过图解和代码掌握”。Planner 智能体先结合画像和薄弱点拆解任务，Supervisor 为本次任务建立唯一运行编号，并调度不同角色协作。讲义智能体负责分层解释，题库智能体生成针对性练习，思维导图智能体整理知识关系，代码智能体创建可运行案例，可视化智能体生成动态实验，拓展阅读和学习路径智能体补充后续内容。各角色共享任务摘要和必要上下文，既保持分工，也避免生成彼此矛盾的材料。生成过程中，界面持续显示规划、检索、生成、审查和保存进度；失败时会显示原因和重试状态，学生不会面对长时间白屏。

【2:35—3:25】
所有资源不会生成后直接发布。Critic 智能体会逐项检查事实依据、结构完整性、难度适配、引用和安全性；发现问题时给出明确返修意见，由对应智能体重新生成。返修次数达到上限仍不合格，就进入人工审核，而不是把低质量内容标记成成功。最终通过 Review 质量门禁后，资源进入个人资源库。这里可以看到讲义、题库、思维导图、学习路径、代码实操、拓展阅读、视频脚本和交互可视化八类资源，每一类都使用适合自己的卡片和交互方式展示。打开任一资源，还能看到生成状态、适配理由和引用依据，方便学生判断是否可信、是否适合当前阶段。

【3:25—4:20】
现在打开 CNN 互动实验。学生可以修改输入矩阵、卷积核、padding 和 stride，观察滑动窗口每一步的覆盖区域以及输出矩阵如何变化。系统不仅给出结果，还会解释输出尺寸公式中每个参数的作用。当参数组合不合法时，页面会直接标明冲突位置和修正建议。随后进入代码实验，学生可以修改 Python 代码、运行示例并编写测试，用计算结果验证刚才的可视化结论。代码运行与后端服务隔离，并设置时间和资源限制。这样就把“看懂概念、观察过程、动手实现、验证结果”连接成一个完整实践任务，而不是只阅读一段静态答案。

【4:20—5:15】
资源生成完成后，系统根据画像、当前掌握度和资源使用记录规划动态学习路径。路径不是固定目录，而是带有前置依赖、预计时长和掌握门槛的有向图。如果学生在卷积输出尺寸练习中连续出错，后续节点会暂时等待，系统先推送公式拆解、基础练习和可视化实验；当证据表明学生已经掌握，路径再继续推进。学生也可以根据自己的时间调整学习节奏，系统会重新计算节点顺序与预计完成时间。推荐卡片会说明推荐原因和资源来源，让学生知道系统为什么此时推送这项内容。

【5:15—6:05】
在学习效果页面，可以看到固定测验、生成练习、错题复习、学习时长、知识点正确率和前后测变化。系统根据这些证据生成学习状态报告，并把结果反馈给画像、路径和下一轮资源生成。报告不只给出一个总分，还会区分知识理解、迁移应用和实践完成情况，并给出下一步建议。学生在学习过程中也可以随时打开辅导面板进行连续提问，回答以流式 Markdown 展示，支持公式、引用和语音。证据不足时，系统会明确说明边界，不会用没有依据的内容填补答案。

【6:05—6:40】
为了降低幻觉，系统采用向量检索与 BM25 融合，从课程知识库中获取依据，并保存引用和验证记录。检索内容会经过提示注入隔离、引用校验和事实检查，授权状态不明确的资料默认不进入提交知识库。Chroma 不可用时自动切换 BM25；模型超时或密钥缺失时使用可解释的降级资源。每次生成还记录模型、Prompt 版本、智能体步骤、耗时、返修和降级原因，便于查看完整审计链路。

【6:40—7:00】
VisionaryTutor 的价值不只是生成更多内容，而是把对话画像、多智能体协作、多模态资源、动态路径和学习评估连接成持续优化的闭环。它让学生获得真正适合当前状态的学习材料，也让教师保留内容审核和教学决策。以上是我们的项目演示，谢谢各位老师。
"""
    doc = Document()
    configure_document(doc, "7 分钟演示视频讲稿")
    add_cover(
        doc,
        "演示材料",
        "7 分钟演示视频讲稿",
        "纯讲解文字版",
        ("文档版本：V1.0", "编制日期：2026 年 7 月 18 日", "适用项目：VisionaryTutor / 智眸学伴"),
    )

    blocks = [block.strip() for block in script.strip().split("\n\n") if block.strip()]
    for index, block in enumerate(blocks, start=1):
        lines = block.splitlines()
        time_range = lines[0].strip("【】")
        narration = "".join(line.strip() for line in lines[1:])
        doc.add_heading(f"{index:02d}  {time_range}", level=1)
        paragraph = add_para(doc, narration, size=11.2, after=10)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.line_spacing = 1.333

    output_path = OUT / "03_7分钟演示视频讲稿.docx"
    doc.save(output_path)
    legacy_path = OUT / "03_7分钟演示视频讲稿.txt"
    if legacy_path.exists():
        legacy_path.unlink()


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="visionary-docs-") as temp:
        build_project_document(Path(temp))
    build_ai_coding_document()
    build_video_script_document()
    for path in sorted(OUT.iterdir()):
        if path.is_file():
            print(f"created: {path.name} ({path.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
